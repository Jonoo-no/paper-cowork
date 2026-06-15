#!/usr/bin/env python3
"""
paper-cowork docx_converter.py

将 Markdown 章节 + 参考文献数据合成为格式化的 .docx 文件。
读取学校模板的样式定义，继承其字体、字号、页面设置等格式。

用法：
  python docx_converter.py \\
    --template 学校模板.docx \\
    --chapters ./05-初稿/ \\
    --references 参考文献.md \\
    --output 论文终稿.docx

依赖：pip install python-docx
"""

import argparse
import json
import os
import re
import sys

try:
    from docx import Document
    from docx.shared import Pt, Cm, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
except ImportError:
    print("错误：需要 python-docx 库。请运行：pip install python-docx")
    sys.exit(1)


# ─── 辅助函数 ─────────────────────────────────────────────

def parse_markdown_frontmatter(text: str) -> dict:
    """解析 Markdown 文件的 frontmatter (--- 包围的 YAML 风格头信息)。"""
    meta = {}
    m = re.match(r'^---\s*\n(.*?)\n---\s*\n', text, re.DOTALL)
    if m:
        for line in m.group(1).strip().split('\n'):
            if ':' in line:
                key, _, val = line.partition(':')
                meta[key.strip()] = val.strip()
        text = text[m.end():]
    return meta, text.strip()


def parse_references(ref_path: str) -> list:
    """从参考文献 Markdown 文件中提取文献条目。"""
    if not os.path.exists(ref_path):
        print(f"⚠️ 参考文献文件不存在: {ref_path}")
        return []

    with open(ref_path, 'r', encoding='utf-8') as f:
        content = f.read()

    refs = []
    # 标准 GB/T 7714 格式行: [1] 作者. 标题[J]...
    pattern = re.compile(r'^\[(\d+)\]\s*(.*)', re.MULTILINE)
    for m in pattern.finditer(content):
        refs.append({
            'index': int(m.group(1)),
            'text': m.group(2).strip()
        })

    if not refs:
        # 尝试解析结构化 frontmatter 格式
        metas, _ = parse_markdown_frontmatter(content)
        if metas:
            refs = [{'index': 1, 'text': metas.get('title', ''), 'meta': metas}]

    return refs


def read_chapters(chapters_dir: str) -> list:
    """读取章节目录下的所有 Markdown 文件，按文件名排序返回。"""
    if not os.path.isdir(chapters_dir):
        print(f"错误：章节目录不存在: {chapters_dir}")
        return []

    md_files = sorted([
        f for f in os.listdir(chapters_dir)
        if f.endswith('.md') and f != '参考文献列表.md'
    ])

    chapters = []
    for fname in md_files:
        fpath = os.path.join(chapters_dir, fname)
        with open(fpath, 'r', encoding='utf-8') as f:
            text = f.read()
        _, body = parse_markdown_frontmatter(text)
        chapters.append({
            'file': fname,
            'path': fpath,
            'body': body.strip()
        })

    return chapters


def parse_markdown_to_paragraphs(md_text: str) -> list:
    """
    将 Markdown 文本解析为段落块列表。
    每块包含类型（heading1/heading2/heading3/normal/list/code/image/table）和内容。
    """
    blocks = []
    lines = md_text.split('\n')
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # 标题
        hm = re.match(r'^(#{1,3})\s+(.+)$', stripped)
        if hm:
            level = len(hm.group(1))
            blocks.append({
                'type': f'heading{level}',
                'text': hm.group(2).strip()
            })
            i += 1
            continue

        # 图片 ![](path)
        imgm = re.match(r'!\[(.*?)\]\((.+?)\)', stripped)
        if imgm:
            blocks.append({
                'type': 'image',
                'alt': imgm.group(1),
                'src': imgm.group(2)
            })
            i += 1
            continue

        # 代码块 (```)
        if stripped.startswith('```'):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                code_lines.append(lines[i])
                i += 1
            i += 1  # 跳过结束 ```
            blocks.append({
                'type': 'code',
                'text': '\n'.join(code_lines)
            })
            continue

        # 无序列表
        if re.match(r'^[-*+]\s+', stripped):
            list_items = []
            while i < len(lines) and re.match(r'^[-*+]\s+', lines[i].strip()):
                item_text = re.sub(r'^[-*+]\s+', '', lines[i].strip())
                list_items.append(item_text)
                i += 1
            blocks.append({
                'type': 'list',
                'items': list_items
            })
            continue

        # 有序列表
        if re.match(r'^\d+[\.\)]\s+', stripped):
            list_items = []
            while i < len(lines) and re.match(r'^\d+[\.\)]\s+', lines[i].strip()):
                item_text = re.sub(r'^\d+[\.\)]\s+', '', lines[i].strip())
                list_items.append(item_text)
                i += 1
            blocks.append({
                'type': 'ordered_list',
                'items': list_items
            })
            continue

        # 表格（简单处理：以 | 开头的行为表格）
        if stripped.startswith('|') and i + 1 < len(lines) and lines[i+1].strip().startswith('|---'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                if not lines[i].strip().startswith('|---'):
                    cells = [c.strip() for c in lines[i].strip().split('|')[1:-1]]
                    rows.append(cells)
                i += 1
            if rows:
                blocks.append({
                    'type': 'table',
                    'rows': rows
                })
            continue

        # 空行
        if stripped == '':
            i += 1
            continue

        # 普通段落（合并连续文本行）
        para_lines = []
        while i < len(lines):
            s = lines[i].strip()
            if s == '' or s.startswith('#') or s.startswith('```') or s.startswith('|'):
                break
            if re.match(r'^[-*\d]', s) and not re.match(r'^[-*+]\s+', s) and not re.match(r'^\d+[\.\)]\s+', s):
                # 但需要检查是否真的是列表
                if re.match(r'^[-*+]\s+', s) or re.match(r'^\d+[\.\)]\s+', s):
                    break
            # 图片行单独处理
            if re.match(r'!\[.*?\]\(.+?\)', s):
                if para_lines:
                    break
                else:
                    blocks.append({
                        'type': 'image',
                        'alt': re.match(r'!\[(.*?)\]\(', s).group(1),
                        'src': re.match(r'!\[.*?\]\((.+?)\)', s).group(1)
                    })
                    i += 1
                    para_lines = None
                    break
            para_lines.append(s)
            i += 1

        if para_lines:
            blocks.append({
                'type': 'normal',
                'text': '\n'.join(para_lines)
            })
            continue

        i += 1

    return blocks


def apply_style(paragraph, style_name: str, template_styles: dict):
    """应用模板样式到段落。"""
    if style_name in template_styles:
        style = template_styles[style_name]
        paragraph.style = paragraph.document.styles[style_name]

        # 复制字体属性
        if 'font' in style:
            for run in paragraph.runs:
                if 'name' in style['font']:
                    run.font.name = style['font']['name']
                    # 中文字体
                    r = run._element
                    rPr = r.find(qn('w:rPr'))
                    if rPr is None:
                        rPr = r.makeelement(qn('w:rPr'), {})
                        r.insert(0, rPr)
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = rPr.makeelement(qn('w:rFonts'), {})
                        rPr.append(rFonts)
                    rFonts.set(qn('w:eastAsia'), style['font'].get('eastAsia', '宋体'))
                if 'size' in style['font']:
                    run.font.size = Pt(style['font']['size'])

        # 复制段落格式
        if 'paragraph' in style:
            pf = paragraph.paragraph_format
            if 'alignment' in style['paragraph']:
                pf.alignment = style['paragraph']['alignment']
            if 'first_line_indent' in style['paragraph']:
                pf.first_line_indent = Cm(style['paragraph']['first_line_indent'])
    else:
        paragraph.style = paragraph.document.styles[style_name]


def extract_template_styles(template_path: str) -> dict:
    """从模板中提取样式定义。"""
    styles = {}
    try:
        doc = Document(template_path)
        for style in doc.styles:
            style_info = {}
            # 字体
            font_info = {}
            if style.font.name:
                font_info['name'] = style.font.name
            if style.font.size:
                font_info['size'] = style.font.size.pt
            if font_info:
                style_info['font'] = font_info

            # 段落格式
            para_info = {}
            pf = style.paragraph_format
            if pf.alignment is not None:
                para_info['alignment'] = pf.alignment
            if pf.first_line_indent:
                para_info['first_line_indent'] = pf.first_line_indent.pt / 28.35  # pt to cm

            if para_info:
                style_info['paragraph'] = para_info

            if style_info:
                styles[style.name] = style_info

        doc.close()
    except Exception as e:
        print(f"⚠️ 读取模板样式时出错: {e}")

    return styles


def build_docx(template_path: str, chapters: list, refs: list, output_path: str):
    """构建格式化的 .docx 文档。"""
    # 检查模板
    if not os.path.exists(template_path):
        # 无模板时创建空白文档
        print("⚠️ 未找到模板文件，使用默认空白文档。")
        doc = Document()
        # 设置默认中文字体
        style = doc.styles['Normal']
        style.font.name = '宋体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        template_styles = {}
    else:
        doc = Document(template_path)
        template_styles = extract_template_styles(template_path)

    # 清空模板中的内容（但保留样式）
    # python-docx 没有直接清空的方法，我们保留模板的结构
    # 如果模板有内容，在写入新内容之前需要清空
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = ''

    # 写入章节内容
    for i, chapter in enumerate(chapters):
        blocks = parse_markdown_to_paragraphs(chapter['body'])

        # 如果第一章文件名包含"绪论"或"引言"等，但用户可能已经写了标题
        # 添加章节分隔标记
        if i > 0 and blocks:
            doc.add_paragraph()  # 章节间空行

        for block in blocks:
            if block['type'] == 'heading1':
                p = doc.add_heading(block['text'], level=1)
                apply_style(p, 'Heading 1', template_styles)

            elif block['type'] == 'heading2':
                p = doc.add_heading(block['text'], level=2)
                apply_style(p, 'Heading 2', template_styles)

            elif block['type'] == 'heading3':
                p = doc.add_heading(block['text'], level=3)
                apply_style(p, 'Heading 3', template_styles)

            elif block['type'] == 'normal':
                p = doc.add_paragraph(block['text'])
                apply_style(p, 'Normal', template_styles)

            elif block['type'] == 'list':
                for item in block['items']:
                    p = doc.add_paragraph(item, style='List Bullet')
                    apply_style(p, 'List Bullet', template_styles)

            elif block['type'] == 'ordered_list':
                for idx, item in enumerate(block['items'], 1):
                    p = doc.add_paragraph(item, style='List Number')
                    apply_style(p, 'List Number', template_styles)

            elif block['type'] == 'code':
                p = doc.add_paragraph(block['text'])
                p.style = doc.styles['Normal']
                pf = p.paragraph_format
                pf.left_indent = Cm(1)
                # 设置等宽字体
                for run in p.runs:
                    run.font.name = 'Courier New'
                    run.font.size = Pt(9)

            elif block['type'] == 'image':
                src = block['src']
                if os.path.exists(src):
                    try:
                        doc.add_picture(src, width=Inches(5))
                        # 图片居中
                        last_paragraph = doc.paragraphs[-1]
                        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        # 如果有 alt 文字，添加为图注
                        if block.get('alt'):
                            caption = doc.add_paragraph(f"图: {block['alt']}")
                            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            for run in caption.runs:
                                run.font.size = Pt(9)
                    except Exception as e:
                        print(f"⚠️ 插入图片失败 {src}: {e}")
                        doc.add_paragraph(f"[图片: {block.get('alt', '未命名')}]")
                else:
                    print(f"⚠️ 图片文件不存在: {src}")
                    doc.add_paragraph(f"[图片文件未找到: {block.get('alt', src)}]")

            elif block['type'] == 'table':
                if block['rows']:
                    table = doc.add_table(rows=len(block['rows']), cols=len(block['rows'][0]))
                    table.style = 'Table Grid'
                    for r_idx, row_data in enumerate(block['rows']):
                        for c_idx, cell_text in enumerate(row_data):
                            if c_idx < len(table.columns):
                                table.cell(r_idx, c_idx).text = cell_text

    # 添加参考文献章节
    if refs:
        doc.add_page_break()
        p = doc.add_heading('参考文献', level=1)
        apply_style(p, 'Heading 1', template_styles)

        for ref in refs:
            ref_text = f"[{ref['index']}] {ref['text']}"
            p = doc.add_paragraph(ref_text)
            apply_style(p, 'Normal', template_styles)
            # 悬挂缩进
            pf = p.paragraph_format
            pf.first_line_indent = Cm(-0.74)
            pf.left_indent = Cm(0.74)

    # 保存
    doc.save(output_path)
    print(f"✅ 论文已生成: {output_path}")


# ─── 主入口 ─────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description='paper-cowork: 将 Markdown 章节转换为格式化的 .docx 论文'
    )
    parser.add_argument('--template', '-t', required=True,
                        help='学校论文模板 .docx 路径')
    parser.add_argument('--chapters', '-c', required=True,
                        help='章节 Markdown 文件所在目录')
    parser.add_argument('--references', '-r',
                        help='参考文献 Markdown 文件路径')
    parser.add_argument('--output', '-o', default='论文终稿.docx',
                        help='输出 .docx 路径 (默认: 论文终稿.docx)')

    args = parser.parse_args()

    # 读取章节
    print(f"📖 读取章节: {args.chapters}")
    chapters = read_chapters(args.chapters)
    print(f"   → 找到 {len(chapters)} 个章节文件")

    if not chapters:
        print("错误: 未找到章节文件。")
        sys.exit(1)

    # 读取参考文献
    refs = []
    if args.references:
        print(f"📚 读取参考文献: {args.references}")
        refs = parse_references(args.references)
        print(f"   → 找到 {len(refs)} 条参考文献")

    # 构建 docx
    build_docx(args.template, chapters, refs, args.output)


if __name__ == '__main__':
    main()
